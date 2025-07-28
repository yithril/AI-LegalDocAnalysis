import aiofiles
import time
from typing import AsyncIterator
from pathlib import Path
from docx import Document
from models.tenant.document import DocumentStatus

from ..text_extraction_strategy import TextExtractionStrategy
from ...models.extraction_result import TextExtractionResult


class WordDocumentStrategy(TextExtractionStrategy):
    async def can_handle(self, file_path: str, mime_type: str) -> bool:
        path = Path(file_path)
        return (
            path.suffix.lower() in self.get_supported_extensions() or
            mime_type in self.get_supported_mime_types()
        )

    async def validate_file(self, file_path: str) -> bool:
        try:
            path = Path(file_path)

            if not path.exists():
                return False
            
            if not path.is_file():
                return False
            
            if await self.is_empty_file(file_path):
                return False
            
            # Try to open and validate as DOCX
            try:
                doc = Document(file_path)
                # If we can open it, it's valid
                return True
            except Exception:
                return False
                
        except Exception:
            return False

    async def extract_text_from_stream(self, file_path: str) -> TextExtractionResult:
        start_time = time.time()
        
        try:
            if not await self.validate_file(file_path):
                processing_time = time.time() - start_time
                return TextExtractionResult.failure_result(
                    status=DocumentStatus.CORRUPTED,
                    file_path=file_path,
                    strategy_used=self.__class__.__name__,
                    error_message=f"Invalid file or corrupted DOCX file: {file_path}",
                    processing_time=processing_time
                )

            async def text_chunks():
                try:
                    doc = Document(file_path)
                    chunk_size = 8192  # Same as other strategies
                    
                    # Extract text from all paragraphs
                    full_text = ""
                    for paragraph in doc.paragraphs:
                        full_text += paragraph.text + "\n"
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                full_text += cell.text + "\t"
                            full_text += "\n"
                    
                    # Split into chunks
                    for i in range(0, len(full_text), chunk_size):
                        chunk = full_text[i:i + chunk_size]
                        if chunk:
                            yield chunk
                            
                except Exception as e:
                    # If text extraction fails, yield error message
                    yield f"Error extracting text from DOCX: {str(e)}"

            processing_time = time.time() - start_time
            return TextExtractionResult.success_result(
                text_chunks=text_chunks(),
                file_path=file_path,
                strategy_used=self.__class__.__name__,
                processing_time=processing_time,
                metadata={
                    "file_size": Path(file_path).stat().st_size,
                    "format": "docx"
                }
            )
        except Exception as e:
            processing_time = time.time() - start_time
            return TextExtractionResult.failure_result(
                status=DocumentStatus.EXTRACTION_FAILED,
                file_path=file_path,
                strategy_used=self.__class__.__name__,
                error_message=str(e),
                processing_time=processing_time
            )

    def get_supported_extensions(self) -> list[str]:
        return ['.docx']

    def get_supported_mime_types(self) -> list[str]:
        return ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'] 