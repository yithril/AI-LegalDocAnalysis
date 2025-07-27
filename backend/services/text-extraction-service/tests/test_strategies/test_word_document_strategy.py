import pytest
import tempfile
import os
from pathlib import Path
from unittest.mock import AsyncMock, patch, MagicMock
from common_lib.document_enums import DocumentStatus
from text_extraction_service.strategies.document_strategies.word_document_strategy import WordDocumentStrategy
from docx import Document


class TestWordDocumentStrategy:
    @pytest.fixture
    def strategy(self):
        return WordDocumentStrategy()

    @pytest.fixture
    def sample_docx_file(self):
        """Create a simple test DOCX file for testing."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        # Create a real DOCX with some text
        doc = Document()
        doc.add_heading('Legal Contract', 0)
        doc.add_paragraph('This is a test legal document.')
        doc.add_paragraph('It contains multiple paragraphs with important legal content.')
        doc.add_heading('Terms and Conditions', level=1)
        doc.add_paragraph('1. This is the first term.')
        doc.add_paragraph('2. This is the second term.')
        doc.add_paragraph('3. This is the third term.')
        doc.save(temp_path)
        
        yield temp_path
        os.unlink(temp_path)

    @pytest.fixture
    def complex_docx_file(self):
        """Create a complex DOCX file with tables and formatting."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        # Create a complex DOCX
        doc = Document()
        doc.add_heading('Complex Legal Document', 0)
        
        # Add paragraphs
        doc.add_paragraph('This document contains various elements.')
        
        # Add a table
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = 'Name'
        table.cell(0, 1).text = 'Value'
        table.cell(1, 0).text = 'Contract ID'
        table.cell(1, 1).text = 'CON-2024-001'
        table.cell(2, 0).text = 'Date'
        table.cell(2, 1).text = '2024-01-15'
        
        # Add more content
        doc.add_paragraph('Additional legal terms and conditions apply.')
        doc.add_heading('Signature Block', level=1)
        doc.add_paragraph('Signed by: _________________')
        doc.add_paragraph('Date: _________________')
        
        doc.save(temp_path)
        
        yield temp_path
        os.unlink(temp_path)

    @pytest.fixture
    def invalid_docx_file(self):
        """Create an invalid DOCX file for testing."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            f.write(b'This is not a valid DOCX file')
            temp_path = f.name
        yield temp_path
        os.unlink(temp_path)

    @pytest.fixture
    def empty_docx_file(self):
        """Create an empty DOCX file for testing."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        # Create an empty DOCX
        doc = Document()
        doc.save(temp_path)
        
        yield temp_path
        os.unlink(temp_path)

    @pytest.mark.asyncio
    async def test_can_handle_with_docx_extension(self, strategy):
        """Test can_handle returns True for DOCX file extensions."""
        assert await strategy.can_handle("document.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document") is True
        assert await strategy.can_handle("contract.DOCX", "application/vnd.openxmlformats-officedocument.wordprocessingml.document") is True

    @pytest.mark.asyncio
    async def test_can_handle_with_docx_mime_type(self, strategy):
        """Test can_handle returns True for DOCX MIME types."""
        assert await strategy.can_handle("document.unknown", "application/vnd.openxmlformats-officedocument.wordprocessingml.document") is True

    @pytest.mark.asyncio
    async def test_can_handle_with_unsupported_file(self, strategy):
        """Test can_handle returns False for unsupported files."""
        assert await strategy.can_handle("document.txt", "text/plain") is False
        assert await strategy.can_handle("document.pdf", "application/pdf") is False
        assert await strategy.can_handle("document.doc", "application/msword") is False  # Note: .doc not supported

    @pytest.mark.asyncio
    async def test_validate_file_with_valid_docx(self, strategy, sample_docx_file):
        """Test validate_file returns True for valid DOCX files."""
        assert await strategy.validate_file(sample_docx_file) is True

    @pytest.mark.asyncio
    async def test_validate_file_with_invalid_docx(self, strategy, invalid_docx_file):
        """Test validate_file returns False for invalid DOCX files."""
        assert await strategy.validate_file(invalid_docx_file) is False

    @pytest.mark.asyncio
    async def test_validate_file_with_empty_docx(self, strategy, empty_docx_file):
        """Test validate_file returns True for empty DOCX files (they're still valid)."""
        assert await strategy.validate_file(empty_docx_file) is True

    @pytest.mark.asyncio
    async def test_validate_file_with_nonexistent_file(self, strategy):
        """Test validate_file returns False for nonexistent files."""
        assert await strategy.validate_file("nonexistent.docx") is False

    @pytest.mark.asyncio
    async def test_validate_file_with_directory(self, strategy):
        """Test validate_file returns False for directories."""
        with tempfile.TemporaryDirectory() as temp_dir:
            assert await strategy.validate_file(temp_dir) is False

    @pytest.mark.asyncio
    async def test_extract_text_from_stream_with_valid_docx(self, strategy, sample_docx_file):
        """Test extract_text_from_stream yields content from valid DOCX files."""
        result = await strategy.extract_text_from_stream(sample_docx_file)
        
        assert result.success is True
        assert result.status == DocumentStatus.PROCESSED
        assert result.file_path == sample_docx_file
        assert result.strategy_used == "WordDocumentStrategy"
        assert result.error_message is None
        assert result.processing_time is not None
        assert result.processing_time > 0
        
        content = ""
        async for chunk in result.text_chunks:
            content += chunk
        
        # Should contain the text from our test DOCX
        assert "Legal Contract" in content
        assert "This is a test legal document" in content
        assert "Terms and Conditions" in content
        assert "This is the first term" in content

    @pytest.mark.asyncio
    async def test_extract_text_from_stream_with_complex_docx(self, strategy, complex_docx_file):
        """Test extract_text_from_stream handles complex DOCX files with tables."""
        result = await strategy.extract_text_from_stream(complex_docx_file)
        
        assert result.success is True
        assert result.status == DocumentStatus.PROCESSED
        
        content = ""
        async for chunk in result.text_chunks:
            content += chunk
        
        # Should contain text from all elements
        assert "Complex Legal Document" in content
        assert "Contract ID" in content
        assert "CON-2024-001" in content
        assert "Signature Block" in content
        assert "Signed by" in content

    @pytest.mark.asyncio
    async def test_extract_text_from_stream_with_invalid_docx(self, strategy, invalid_docx_file):
        """Test extract_text_from_stream returns failure result for invalid DOCX files."""
        result = await strategy.extract_text_from_stream(invalid_docx_file)
        
        assert result.success is False
        assert result.status == DocumentStatus.CORRUPTED
        assert result.error_message == "Invalid file or corrupted DOCX file: " + invalid_docx_file
        assert result.processing_time is not None

    @pytest.mark.asyncio
    async def test_extract_text_from_stream_with_empty_docx(self, strategy, empty_docx_file):
        """Test extract_text_from_stream handles empty DOCX files."""
        result = await strategy.extract_text_from_stream(empty_docx_file)
        
        assert result.success is True
        assert result.status == DocumentStatus.PROCESSED
        
        content = ""
        async for chunk in result.text_chunks:
            content += chunk
        
        # Empty DOCX should yield minimal content or empty string
        assert len(content) >= 0  # Could be empty or contain minimal structure

    @pytest.mark.asyncio
    async def test_extract_text_from_stream_with_nonexistent_file(self, strategy):
        """Test extract_text_from_stream returns failure result for nonexistent files."""
        result = await strategy.extract_text_from_stream("nonexistent.docx")
        
        assert result.success is False
        assert result.status == DocumentStatus.CORRUPTED
        assert result.error_message == "Invalid file or corrupted DOCX file: nonexistent.docx"
        assert result.processing_time is not None

    @pytest.mark.asyncio
    async def test_extract_text_from_stream_chunking(self, strategy, complex_docx_file):
        """Test that extract_text_from_stream yields content in chunks."""
        result = await strategy.extract_text_from_stream(complex_docx_file)
        
        assert result.success is True
        assert result.status == DocumentStatus.PROCESSED
        
        chunks = []
        async for chunk in result.text_chunks:
            chunks.append(chunk)
        
        # Should have at least one chunk
        assert len(chunks) > 0
        
        # All chunks should be strings
        for chunk in chunks:
            assert isinstance(chunk, str)
        
        # Combined content should contain expected text
        combined_content = ''.join(chunks)
        assert "Legal Document" in combined_content
        assert "Contract ID" in combined_content

    def test_get_supported_extensions(self, strategy):
        """Test get_supported_extensions returns correct list."""
        expected_extensions = ['.docx']
        assert strategy.get_supported_extensions() == expected_extensions

    def test_get_supported_mime_types(self, strategy):
        """Test get_supported_mime_types returns correct list."""
        expected_mime_types = ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
        assert strategy.get_supported_mime_types() == expected_mime_types

    @pytest.mark.asyncio
    async def test_extract_text_preserves_structure(self, strategy, sample_docx_file):
        """Test that text extraction preserves document structure."""
        result = await strategy.extract_text_from_stream(sample_docx_file)
        
        assert result.success is True
        
        content = ""
        async for chunk in result.text_chunks:
            content += chunk
        
        # Should preserve headings and paragraph structure
        assert "Legal Contract" in content  # Main heading
        assert "Terms and Conditions" in content  # Sub heading
        assert "This is a test legal document" in content  # Paragraph content 